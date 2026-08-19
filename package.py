#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 Windows++ v3.0 使用说明 Word 文档，并打包 zip（含源码、wpp 包、exe）到桌面。"""
import os
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
VERSION = "4.0"

import ctypes
_buf = ctypes.create_unicode_buffer(260)
ctypes.windll.shell32.SHGetFolderPathW(None, 0, None, 0, _buf)
DESKTOP = _buf.value or os.path.join(os.path.expanduser("~"), "Desktop")

HEADER = RGBColor(0x1F, 0x29, 0x37)


def h(doc, text, size=14, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = HEADER
    return p


def para(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def code_block(doc, text, size=9.5):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x0B, 0x53, 0x4F)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def feature_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (a, b) in enumerate(rows):
        c0 = table.cell(i, 0)
        c0.text = ""
        r = c0.paragraphs[0].add_run(a)
        r.bold = True
        r.font.size = Pt(10)
        c1 = table.cell(i, 1)
        c1.text = ""
        r = c1.paragraphs[0].add_run(b)
        r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx(path):
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Windows++ 使用说明")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = HEADER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"软件管家 v{VERSION}（更新 / 卸载 / 清理 / 桌面锁定 / 个性化 / 工具宠物）  |  更新日期：2026-08-19")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    h(doc, "一、软件简介")
    para(doc, "Windows++ 是一款一站式 Windows 软件管家，主窗口左侧导航栏集成六大功能模块：")
    feature_table(doc, [
        ("⬆ 软件更新", "注册表枚举 + winget 比对，一键批量更新，支持暂停/跳过/取消、实时状态、右键取消勾选、导出报告"),
        ("🗑 软件卸载", "已装软件列表，名称搜索、按名称/大小/日期排序，调用系统卸载程序并反馈结果"),
        ("🧹 扫描清理", "全盘扫描旧版本残留文件 + 下载未清理的安装包，进度条、勾选列表、全选/全不选、清理选中"),
        ("🔒 桌面锁定", "一键扫描桌面图标、逐图标锁定（持久化）；被锁图标被移动/删除时弹窗提示并可一键前往解锁"),
        ("⚙ 设置", "Windows 10/11 色板主题换肤；背景图片/视频（不透明度滑块默认 35%）；背景音乐；开机自启动开关"),
        ("🧸 工具与宠物", "联网下载安装 Win7 经典桌面小工具（时钟/日历/CPU 仪表盘/天气等）；桌面宠物（默认 FeibiPet）"),
    ])

    h(doc, "二、环境要求")
    bullet(doc, "Windows 10（1809 及以上）或 Windows 11；")
    bullet(doc, "winget（Windows 包管理器）：Win11 与新版 Win10 通常自带；缺失时在微软商店安装「应用安装程序」或访问 https://aka.ms/getwinget ；")
    bullet(doc, "源码运行需要 Python 3.7+（含 tkinter）；也可直接使用打包好的 WindowsPP.exe（免安装）。")
    para(doc, "源码版为纯标准库实现，无需安装任何第三方 Python 包。", bold=True)

    h(doc, "三、启动方式")
    bullet(doc, "方式一（推荐）：双击 WindowsPP.exe（已编译版，免安装）或桌面快捷方式；")
    bullet(doc, "方式二：双击「启动WindowsPP.bat」自动探测 Python 解释器；")
    bullet(doc, "方式三（命令行）：", size=10.5)
    code_block(doc, "python WindowsPP.py           # 打开图形界面\n"
                    "python WindowsPP.py --scan    # 只读扫描（不执行更新）\n"
                    "python WindowsPP.py --tray   # 后台最小化运行（开机启动场景）")

    h(doc, "四、界面总览")
    para(doc, "主窗口左侧为导航栏，自上而下排列：软件更新（默认选中）、软件卸载、扫描清理、桌面锁定、设置、工具与宠物。"
              "点击切换右侧功能面板，底部状态栏统一显示当前进度与提示。")

    h(doc, "五、软件更新")
    para(doc, "1. 点击「🔄 扫描并检查更新」：列出全部已装软件，可更新项默认勾选，注册表版本已达标者显示绿色「已是最新」；")
    para(doc, "2. 点击行首 ✔ 列或右键可取消/恢复勾选；")
    para(doc, "3. 点击「⬆ 开始更新」：确认清单（可勾选更新后自动清理旧文件/安装包），逐个通过 winget 升级；")
    para(doc, "4. 更新过程可随时「⏸ 暂停 / ⏭ 跳过当前 / ✖ 取消当前 / ⏹ 取消全部」，状态列实时反馈；")
    para(doc, "5. 支持「🔍 扫描旧版文件」「🧹 清理旧版本文件」「🗑 清除安装包」「📄 导出报告」。")

    h(doc, "六、软件卸载")
    para(doc, "列表展示全部已安装软件，顶部支持按名称搜索、按名称/大小/安装日期排序；选中后点击「🗑 卸载选中」或双击，"
              "程序调用该软件自带的卸载程序（不静默删除任何文件），完成后自动刷新列表。卸载不可恢复，请谨慎操作。")

    h(doc, "七、扫描清理")
    para(doc, "点击「🔍 开始扫描」：程序全盘扫描两类内容并显示进度条——①软件旧版本残留目录（同目录并存多版本，保留最新，"
              "不触碰登录与使用数据）；②下载/临时目录中的安装包（.exe/.msi 等）。结果以列表呈现，逐项勾选后点击「🗑 清理选中」"
              "确认删除；支持「☑ 全选 / ☐ 全不选」。")

    h(doc, "八、桌面锁定")
    para(doc, "1. 「🔍 一键扫描桌面图标」：列出桌面全部图标及其类型；")
    para(doc, "2. 点击每行的「🔒」列可单独锁定/解锁该图标（状态持久化保存）；")
    para(doc, "3. 锁定后，该图标被移动、删除或重命名时程序会弹窗提示「桌面图标已锁定」，可一键前往解锁；")
    para(doc, "4. 顶部「全局锁定」开关：启用桌面图标自动排列，从系统层面禁止拖乱。")
    para(doc, "说明：Windows 未提供阻止拖拽的官方接口，真正的「禁止移动」需配合全局锁定（自动排列）使用。")

    h(doc, "九、设置")
    bullet(doc, "开机自启动：Windows++ 自身开机后台运行开关（写入注册表「启动」项）；")
    bullet(doc, "颜色主题：Windows 10/11 个性化色板，点击即应用并持久化；")
    bullet(doc, "背景图片：选择图片作为主窗口背景（缩放/透明度需 pillow，默认不透明度 35%）；")
    bullet(doc, "背景视频：选择视频以窗口背景方式播放，默认播放视频音轨，可设静音；")
    bullet(doc, "背景音乐：独立导入音频作为背景音乐（视频未静音时优先视频音轨）。")

    h(doc, "十、工具与宠物")
    para(doc, "工具：内置 Windows 7 经典小工具库（时钟/日历/CPU 仪表盘/天气/货币/资讯/幻灯片/拼图），"
              "点击「📥 安装」联网下载并解压到本地（%LOCALAPPDATA%\\WindowsPP\\gadgets），「▶ 打开」以兼容方式"
              "（默认浏览器显示小工具内容）运行，支持卸载与开机自启动。"
              "提示：Windows 10/11 已移除侧边栏平台，小工具无法像 Win7 那样在桌面侧边栏原生运行。")
    para(doc, "宠物：默认读取 FeibiPet 目录（C:\\Users\\…\\my file\\FeibiPet_v0.0.1），支持指定/导入自定义宠物 exe；"
              "提供「打开/关闭宠物」与「开机自启动」开关。")

    h(doc, "十一、常见问题（FAQ）")
    faqs = [
        ("启动提示找不到 winget？", "在微软商店安装「应用安装程序」（App Installer），装完重开程序。"),
        ("更新失败怎么办？", "看下方日志错误尾巴。常见原因：软件正在运行、需要管理员权限、网络问题。"
         "可稍后手动执行：winget upgrade --id <ID> -e"),
        ("部分软件状态是「—」（灰色）？", "winget 未收录该软件的更新信息，无法自动判定/升级，属正常现象。"),
        ("小工具下载失败？", "检查网络连接后重试；下载源为 Win7 官方小工具存档（GitHub Pages），一般可直接访问。"),
        ("背景图片不显示？", "图片缩放/透明度需要 pillow（pip install pillow）；未安装时显示原尺寸。"),
        ("桌面锁定拦不住拖拽？", "请同时开启「全局锁定（自动排列）」；Windows 无阻止拖拽的官方接口。"),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("Q: " + q)
        r.bold = True
        r.font.size = Pt(10.5)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("A: " + a)
        r.font.size = Pt(10.5)

    h(doc, "十二、安全与免责说明")
    bullet(doc, "更新均调用系统自带 winget，安装包来自 winget 官方源；")
    bullet(doc, "卸载仅调用软件自带卸载程序，不静默删除任何文件；")
    bullet(doc, "清理旧版本只删“同目录并存多版本”的旧程序目录，不触碰登录与使用数据；")
    bullet(doc, "本工具仅做版本检查、更新与辅助管理，请勿用于误操作。")

    doc.save(path)
    return path


def collect_source_files():
    """收集源码与资源文件（含 wpp 包递归）。"""
    files = ["WindowsPP.py", "wpp_tools.py", "WindowsPP.ico",
             "启动WindowsPP.bat", "make_icon.py", "make_shortcut.py",
             "README.md", "LICENSE"]
    wpp_dir = os.path.join(BASE, "wpp")
    for root, dirs, names in os.walk(wpp_dir):
        for n in names:
            if n.endswith(".py"):
                full = os.path.join(root, n)
                rel = os.path.relpath(full, BASE).replace("\\", "/")
                files.append(rel)
    return files


def make_zip(zip_path, docx_path, exe_path=None):
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in collect_source_files():
            z.write(os.path.join(BASE, f), arcname=f)
        z.write(docx_path, arcname=os.path.basename(docx_path))
        if exe_path and os.path.isfile(exe_path):
            z.write(exe_path, arcname=os.path.basename(exe_path))
    return zip_path


def main():
    docx_path = os.path.join(BASE, "Windows++ 使用说明.docx")
    build_docx(docx_path)
    print(f"DOCX OK: {os.path.getsize(docx_path)} bytes")

    exe_path = os.path.join(BASE, "dist", "WindowsPP.exe")
    zip_path = os.path.join(BASE, f"Windows++_v{VERSION}.zip")
    make_zip(zip_path, docx_path, exe_path)
    print(f"ZIP OK: {os.path.getsize(zip_path)} bytes")

    import shutil
    shutil.copy2(docx_path, os.path.join(DESKTOP, os.path.basename(docx_path)))
    shutil.copy2(zip_path, os.path.join(DESKTOP, os.path.basename(zip_path)))
    print(f"DESKTOP: {os.path.basename(zip_path)} exists={os.path.exists(os.path.join(DESKTOP, os.path.basename(zip_path)))}")

    with zipfile.ZipFile(zip_path) as z:
        n = len(z.infolist())
        print(f"ZIP CONTENTS: {n} 项")
        for info in z.infolist()[:6]:
            print(f"  {info.filename}  ({info.file_size} bytes)")
        print("  …")


if __name__ == "__main__":
    main()
